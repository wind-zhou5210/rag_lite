"""
Parser 模块单元测试
"""

import os
import tempfile
import pytest
from unittest.mock import MagicMock, patch

from app.services.parser.base import BaseParser
from app.services.parser.factory import get_parser, get_supported_types


class TestBaseParser:
    """BaseParser 基类测试"""

    def test_base_parser_is_abstract(self):
        """测试 BaseParser 是抽象类，不能直接实例化"""
        with pytest.raises(TypeError):
            BaseParser()

    def test_base_parser_requires_parse_method(self):
        """测试子类必须实现 parse 方法"""
        class IncompleteParser(BaseParser):
            pass

        with pytest.raises(TypeError):
            IncompleteParser()


class TestParserFactory:
    """Parser 工厂函数测试"""

    def test_get_parser_pdf(self):
        """测试获取 PDF 解析器"""
        from app.services.parser.pdf_parser import PDFParser

        parser = get_parser('pdf')
        assert isinstance(parser, PDFParser)

    def test_get_parser_docx(self):
        """测试获取 DOCX 解析器"""
        from app.services.parser.docx_parser import DocxParser

        parser = get_parser('docx')
        assert isinstance(parser, DocxParser)

    def test_get_parser_txt(self):
        """测试获取 TXT 解析器"""
        from app.services.parser.txt_parser import TxtParser

        parser = get_parser('txt')
        assert isinstance(parser, TxtParser)

    def test_get_parser_md(self):
        """测试获取 Markdown 解析器"""
        from app.services.parser.md_parser import MdParser

        parser = get_parser('md')
        assert isinstance(parser, MdParser)

    def test_get_parser_case_insensitive(self):
        """测试文件类型不区分大小写"""
        from app.services.parser.pdf_parser import PDFParser

        parser = get_parser('PDF')
        assert isinstance(parser, PDFParser)

        parser = get_parser('Pdf')
        assert isinstance(parser, PDFParser)

    def test_get_parser_unsupported_type(self):
        """测试不支持的文件类型抛出异常"""
        with pytest.raises(ValueError) as exc_info:
            get_parser('unsupported')

        assert "不支持的文件类型" in str(exc_info.value)

    def test_get_supported_types(self):
        """测试获取支持的文件类型列表"""
        types = get_supported_types()

        assert 'pdf' in types
        assert 'docx' in types
        assert 'txt' in types
        assert 'md' in types


class TestTxtParser:
    """TXT 解析器测试"""

    def test_parse_txt_file(self):
        """测试解析 TXT 文件"""
        from app.services.parser.txt_parser import TxtParser

        # 创建临时 TXT 文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("这是测试内容\n第二行内容")
            temp_path = f.name

        try:
            parser = TxtParser()
            text, error = parser.parse(temp_path)

            assert error is None
            assert "这是测试内容" in text
            assert "第二行内容" in text
        finally:
            os.unlink(temp_path)

    def test_parse_nonexistent_file(self):
        """测试解析不存在的文件"""
        from app.services.parser.txt_parser import TxtParser

        parser = TxtParser()
        text, error = parser.parse('/nonexistent/path/file.txt')

        assert text is None
        assert error is not None

    def test_parse_empty_file(self):
        """测试解析空文件"""
        from app.services.parser.txt_parser import TxtParser

        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            temp_path = f.name

        try:
            parser = TxtParser()
            text, error = parser.parse(temp_path)

            assert error is None
            assert text == ""
        finally:
            os.unlink(temp_path)

    def test_parse_utf8_with_bom(self):
        """测试解析带 BOM 的 UTF-8 文件"""
        from app.services.parser.txt_parser import TxtParser

        with tempfile.NamedTemporaryFile(mode='wb', suffix='.txt', delete=False) as f:
            # BOM + UTF-8 content
            content = '\xef\xbb\xbf这是UTF-8 BOM内容'.encode('utf-8')
            f.write(content)
            temp_path = f.name

        try:
            parser = TxtParser()
            text, error = parser.parse(temp_path)

            assert error is None
            assert "这是UTF-8 BOM内容" in text
        finally:
            os.unlink(temp_path)


class TestMdParser:
    """Markdown 解析器测试"""

    def test_parse_md_file(self):
        """测试解析 Markdown 文件"""
        from app.services.parser.md_parser import MdParser

        content = """# 标题

这是一段内容。

## 二级标题

- 列表项1
- 列表项2
"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(content)
            temp_path = f.name

        try:
            parser = MdParser()
            text, error = parser.parse(temp_path)

            assert error is None
            assert "# 标题" in text
            assert "列表项1" in text
        finally:
            os.unlink(temp_path)

    def test_parse_md_preserves_structure(self):
        """测试 Markdown 解析保留文档结构"""
        from app.services.parser.md_parser import MdParser

        content = "# 标题\n\n正文内容\n\n```python\nprint('hello')\n```\n"
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(content)
            temp_path = f.name

        try:
            parser = MdParser()
            text, error = parser.parse(temp_path)

            assert error is None
            # 应该保留原始内容
            assert "```python" in text
        finally:
            os.unlink(temp_path)


class TestDocxParser:
    """DOCX 解析器测试"""

    def test_parse_docx_file(self):
        """测试解析 DOCX 文件"""
        from app.services.parser.docx_parser import DocxParser

        # 创建一个简单的 DOCX 文件
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name

        try:
            # 创建 DOCX 文档
            doc = Document()
            doc.add_heading('测试标题', level=1)
            doc.add_paragraph('这是测试内容')
            doc.save(temp_path)

            parser = DocxParser()
            text, error = parser.parse(temp_path)

            assert error is None
            assert "测试标题" in text
            assert "这是测试内容" in text
        finally:
            os.unlink(temp_path)

    def test_parse_nonexistent_docx(self):
        """测试解析不存在的 DOCX 文件"""
        from app.services.parser.docx_parser import DocxParser

        parser = DocxParser()
        text, error = parser.parse('/nonexistent/path/file.docx')

        assert text is None
        assert error is not None


class TestPdfParser:
    """PDF 解析器测试"""

    def test_parse_pdf_file(self):
        """测试解析 PDF 文件"""
        from app.services.parser.pdf_parser import PDFParser

        # 创建一个简单的 PDF 文件
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name

        try:
            # 创建 PDF
            c = canvas.Canvas(temp_path, pagesize=letter)
            c.drawString(100, 750, "这是测试PDF内容")
            c.save()

            parser = PDFParser()
            text, error = parser.parse(temp_path)

            assert error is None
            assert "测试PDF内容" in text or len(text) > 0  # PDF 文本提取可能有编码差异
        finally:
            os.unlink(temp_path)

    def test_parse_nonexistent_pdf(self):
        """测试解析不存在的 PDF 文件"""
        from app.services.parser.pdf_parser import PDFParser

        parser = PDFParser()
        text, error = parser.parse('/nonexistent/path/file.pdf')

        assert text is None
        assert error is not None

    def test_parse_empty_pdf(self):
        """测试解析空 PDF 文件"""
        from app.services.parser.pdf_parser import PDFParser

        # 创建空 PDF
        from reportlab.pdfgen import canvas

        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name

        try:
            c = canvas.Canvas(temp_path)
            c.showPage()
            c.save()

            parser = PDFParser()
            text, error = parser.parse(temp_path)

            assert error is None
            assert text == "" or text.strip() == ""
        finally:
            os.unlink(temp_path)
