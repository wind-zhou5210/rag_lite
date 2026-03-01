"""
DOCX 文件解析器
"""

import os
from typing import Tuple, Optional

from app.services.parser.base import BaseParser
from app.util.logger import get_logger

logger = get_logger(__name__)


class DocxParser(BaseParser):
    """DOCX 文件解析器

    使用 python-docx 库解析 Word 文档。
    """

    def parse(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """
        解析 DOCX 文件

        Args:
            file_path: 文件路径

        Returns:
            Tuple[Optional[str], Optional[str]]: (文本内容, 错误信息)
        """
        if not os.path.exists(file_path):
            error = f"文件不存在: {file_path}"
            logger.error(error)
            return None, error

        try:
            from docx import Document

            doc = Document(file_path)

            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))

            content = '\n\n'.join(paragraphs)
            logger.debug(f"DOCX 文件解析成功: {file_path}, 段落数: {len(paragraphs)}")

            return content, None

        except ImportError:
            error = "缺少 python-docx 库，请安装: pip install python-docx"
            logger.error(error)
            return None, error

        except Exception as e:
            error = f"解析 DOCX 文件失败: {str(e)}"
            logger.error(error)
            return None, error
